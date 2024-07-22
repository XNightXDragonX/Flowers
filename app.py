from flask import Flask, abort, render_template, url_for, flash, redirect, request, session, send_file
from flask_migrate import Migrate
from config import Config
from models import db, bcrypt, login_manager, User, Order, Flower
from forms import RegistrationForm, LoginForm
from flask_login import login_user, current_user, logout_user, login_required
from flask_graphql import GraphQLView
from schema import schema
from decorators import admin_required
from getpass import getpass
import sys
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io

app = Flask(__name__)
app.config.from_object(Config)

db.init_app(app)
bcrypt.init_app(app)
login_manager.init_app(app)
migrate = Migrate(app, db)

# Создание таблиц при первом запуске
with app.app_context():
    db.create_all()

# Добавление некоторых цветов в базу данных при первом запуске
with app.app_context():
    if not Flower.query.first():
        flowers = [
            Flower(name='Роза', image_url='images/rose.jpg', length=51, price=150),
            Flower(name='Тюльпан', image_url='images/tulip.jpg', length=62, price=120),
            Flower(name='Лилия', image_url='images/lily.jpg', length=56, price=180)
        ]
        db.session.bulk_save_objects(flowers)
        db.session.commit()

# Главная страница
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if not current_user.is_authenticated:
            flash('Пожалуйста, войдите в систему, чтобы сделать заказ.', 'danger')
            return redirect(url_for('login'))

        name = request.form['name']
        address = request.form['address']
        flower_types = request.form.getlist('flower_type')
        message = request.form['message']

        quantities = {}
        for flower_type in flower_types:
            quantity = request.form.get(f'quantity_{flower_type.lower()}', 0)
            quantities[flower_type] = quantity
        
        flower_type_str = ','.join([f"{flower} ({quantities[flower]} шт.)" for flower in flower_types])

        new_order = Order(name=name, address=address, flower_type=flower_type_str, message=message, user_id=current_user.id)
        db.session.add(new_order)
        db.session.commit()

        session['order_created'] = True  # Установить переменную сессии для указания, что заказ был создан
        return redirect(url_for('index'))

    search_query = request.args.get('search', '')
    length_filter = request.args.get('length', '')
    price_filter = request.args.get('price', '')

    flowers = Flower.query.filter(Flower.name.contains(search_query))

    if length_filter:
        min_length, max_length = map(float, length_filter.split('-'))
        flowers = flowers.filter(Flower.length.between(min_length, max_length))
    if price_filter:
        min_price, max_price = map(float, price_filter.split('-'))
        flowers = flowers.filter(Flower.price.between(min_price, max_price))

    flowers = flowers.all()

    order_created = session.pop('order_created', False)  # Проверить и удалить переменную сессии
    return render_template('index.html', flowers=flowers, search_query=search_query, length_filter=length_filter, price_filter=price_filter, order_created=order_created)

# Страница регистрации
@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = RegistrationForm()
    if form.validate_on_submit():
        email_exists = User.query.filter_by(email=form.email.data).first()
        username_exists = User.query.filter_by(username=form.username.data).first()
        if email_exists:
            flash('Email уже зарегистрирован. Пожалуйста, используйте другой email.', 'danger')
        elif username_exists:
            flash('Имя пользователя уже занято. Пожалуйста, используйте другое имя пользователя.', 'danger')
        else:
            hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
            user = User(username=form.username.data, email=form.email.data, password=hashed_password)
            db.session.add(user)
            db.session.commit()
            flash('Ваша учётная запись создана! Теперь вы можете войти.', 'success')
            return redirect(url_for('login'))
    return render_template('register.html', form=form)

# Страница входа
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            login_user(user, remember=form.remember.data)
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('index'))
        else:
            flash('Неудачный вход. Пожалуйста, проверьте email и пароль.', 'danger')
    return render_template('login.html', form=form)

# Страница профиля
@app.route('/profile')
@login_required
def profile():
    orders = Order.query.filter_by(user_id=current_user.id).all()
    order_numbers = {order.id: idx + 1 for idx, order in enumerate(orders)}
    return render_template('profile.html', user=current_user, orders=orders, order_numbers=order_numbers)

# Маршрут для выхода
@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('login'))

def admin_required(f):
    @login_required
    def decorated_function(*args, **kwargs):
        if current_user.role != 'admin':
            abort(403)
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    return decorated_function

@app.cli.command('create_admin')
def create_admin():
    username = input('Введите имя пользователя: ')
    email = input('Введите email: ')
    password = input('Введите пароль: ')

    if User.query.filter_by(username=username).first():
        print('Пользователь с таким именем уже существует.')
        return
    if User.query.filter_by(email=email).first():
        print('Пользователь с таким email уже существует.')
        return

    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
    admin_user = User(username=username, email=email, password=hashed_password, role='admin')
    db.session.add(admin_user)
    db.session.commit()
    print('Администратор успешно создан.')

app.add_url_rule(
    '/graphql',
    view_func=admin_required(GraphQLView.as_view('graphql', schema=schema, graphiql=True)),
    methods=['GET', 'POST']
)

# Функция для генерации DOCX
def create_docx(order):
    doc = Document()
    doc.add_heading('Заказ #{}'.format(order['order_number']), 0)
    doc.add_paragraph('Имя получателя: {}'.format(order['name']))
    doc.add_paragraph('Адрес: {}'.format(order['address']))
    doc.add_paragraph('Тип цветов: {}'.format(order['flower_type']))
    doc.add_paragraph('Сообщение: {}'.format(order['message']))
    return doc

# Маршрут для скачивания DOCX
@app.route('/download/docx/<int:order_id>')
@login_required
def download_docx(order_id):
    order = Order.query.get_or_404(order_id)
    if order.user_id != current_user.id:
        abort(403)
    
    order_data = {
        'order_number': order.id,
        'name': order.name,
        'address': order.address,
        'flower_type': order.flower_type,
        'message': order.message
    }

    doc = create_docx(order_data)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'order_{order.id}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Функция для генерации PDF
def create_pdf(order):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, 'Заказ #{}'.format(order['order_number']))
    c.drawString(100, 725, 'Имя получателя: {}'.format(order['name']))
    c.drawString(100, 700, 'Адрес: {}'.format(order['address']))
    c.drawString(100, 675, 'Тип цветов: {}'.format(order['flower_type']))
    c.drawString(100, 650, 'Сообщение: {}'.format(order['message']))
    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

# Маршрут для скачивания PDF
@app.route('/download/pdf/<int:order_id>')
@login_required
def download_pdf(order_id):
    order = Order.query.get_or_404(order_id)
    if order.user_id != current_user.id:
        abort(403)
    
    order_data = {
        'order_number': order.id,
        'name': order.name,
        'address': order.address,
        'flower_type': order.flower_type,
        'message': order.message
    }

    pdf_buffer = create_pdf(order_data)
    return send_file(pdf_buffer, as_attachment=True, download_name=f'order_{order.id}.pdf', mimetype='application/pdf')

if __name__ == '__main__':
    app.run(debug=True)